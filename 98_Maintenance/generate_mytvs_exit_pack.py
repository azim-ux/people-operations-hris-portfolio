#!/usr/bin/env python3
"""Generate an editable, manager-reviewable employment exit document pack.

The pack deliberately leaves unverified facts as highlighted placeholders. It
does not create signatures, stamps, backdated appointment letters, or completed
salary records.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "08_Employment_Exit_Documents" / "MyTVS_Exit_Pack_2026"

EMPLOYEE = "Mohammad Azimuddin"
COMPANY = "myTVS"
BRANCH = "Aligarh, Uttar Pradesh"
INTERNSHIP_DATES = "26 May 2025 to 16 July 2025"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(38, 38, 38)
GREY = RGBColor(102, 102, 102)
RED = RGBColor(192, 0, 0)
LIGHT_BLUE = "D9EAF7"
LIGHT_GREY = "E7E6E6"
PLACEHOLDER_RE = re.compile(r"(\[[A-Z0-9][A-Z0-9 /(),.&:+\-–—']+\])")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, bold=False, size=9, color=DARK) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_run_font(run, size=10.5, bold=False, italic=False, color=DARK) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_paragraph(
    doc,
    text="",
    *,
    bold=False,
    italic=False,
    size=10.5,
    align=None,
    before=0,
    after=6,
    line=1.08,
    color=DARK,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    set_run_font(r, size=10)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=10)
    return p


def add_company_header(doc, title: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(COMPANY)
    set_run_font(r, size=14, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(BRANCH)
    set_run_font(r, size=8.5, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title.upper())
    set_run_font(r, size=13, bold=True, color=DARK)


def add_draft_banner(doc) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, "FCE4D6")
    set_cell_text(
        cell,
        "DRAFT — FOR INDEPENDENT MANAGER/AUTHORISED-SIGNATORY REVIEW. "
        "NOT VALID UNTIL FACT-CHECKED, SIGNED AND, WHERE USED, STAMPED.",
        bold=True,
        size=8.5,
        color=RED,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_issue_meta(doc, reference=True) -> None:
    table = doc.add_table(rows=2 if reference else 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(3.3)
    table.columns[1].width = Inches(3.3)
    if reference:
        set_cell_text(table.cell(0, 0), "Reference: [COMPANY REFERENCE NUMBER]", size=9)
        set_cell_text(table.cell(0, 1), "Date of issue: [ACTUAL SIGNING DATE]", size=9)
        set_cell_text(table.cell(1, 0), "Employee ID: [EMPLOYEE ID]", size=9)
        set_cell_text(table.cell(1, 1), "Place: Aligarh, Uttar Pradesh", size=9)
    else:
        set_cell_text(table.cell(0, 0), "Date: [ACTUAL DATE]", size=9)
        set_cell_text(table.cell(0, 1), "Place: Aligarh, Uttar Pradesh", size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_signatory_block(doc, *, employee=False, second=False) -> None:
    if second:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        left, right = table.cell(0, 0), table.cell(0, 1)
        for cell in (left, right):
            cell.text = ""
            set_cell_margins(cell, top=180)
        lp = left.paragraphs[0]
        for text in (
            "\n\n_______________________________\n"
            f"For {COMPANY}\n"
            "Name: [AUTHORISED SIGNATORY NAME]\n"
            "Designation: [SIGNATORY DESIGNATION]\n"
            "Official email/phone: [CONTACT]\n"
            "Company stamp:"
        ):
            r = lp.add_run(text)
            set_run_font(r, size=9.5)
        rp = right.paragraphs[0]
        for text in (
            "\n\n_______________________________\n"
            f"{EMPLOYEE}\n"
            "Employee ID: [EMPLOYEE ID]\n"
            "Date received: [DATE]\n"
            "Signature:"
        ):
            r = rp.add_run(text)
            set_run_font(r, size=9.5)
        return

    add_paragraph(doc, "\n\n_______________________________", after=2)
    if employee:
        add_paragraph(doc, EMPLOYEE, bold=True, after=1)
        add_paragraph(doc, "Employee ID: [EMPLOYEE ID]", after=1)
        add_paragraph(doc, "Mobile/email: [EMPLOYEE CONTACT]", after=1)
        add_paragraph(doc, "Signature: ____________________", after=1)
    else:
        add_paragraph(doc, f"For {COMPANY}", bold=True, after=3)
        add_paragraph(doc, "Name: [AUTHORISED SIGNATORY NAME]", after=1)
        add_paragraph(doc, "Designation: [SIGNATORY DESIGNATION]", after=1)
        add_paragraph(doc, "Official email/phone: [CONTACT]", after=1)
        add_paragraph(doc, "Company stamp:", after=1)


def add_footer(doc) -> None:
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            "Editable draft prepared for factual verification. "
            "Issue only on approved letterhead with an authorised signature."
        )
        set_run_font(r, size=7.5, color=GREY)


def configure_doc(doc, title: str) -> None:
    sec = doc.sections[0]
    sec.top_margin = Cm(1.55)
    sec.bottom_margin = Cm(1.45)
    sec.left_margin = Cm(1.8)
    sec.right_margin = Cm(1.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    styles["Title"].font.name = "Arial"
    styles["Title"].font.color.rgb = BLUE
    styles["Title"].font.size = Pt(17)
    styles["Title"].font.bold = True

    doc.core_properties.title = title
    doc.core_properties.author = EMPLOYEE
    doc.core_properties.subject = "Employment exit document draft for manager review"
    doc.core_properties.comments = (
        "Unverified fields are highlighted. No document is valid until independently "
        "verified and signed by an authorised company representative."
    )


def iter_paragraphs(parent):
    for p in parent.paragraphs:
        yield p
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def highlight_placeholders(doc) -> None:
    for p in iter_paragraphs(doc):
        for run in list(p.runs):
            if not PLACEHOLDER_RE.search(run.text):
                continue
            text = run.text
            props = {
                "bold": run.bold,
                "italic": run.italic,
                "size": run.font.size,
                "name": run.font.name,
                "color": run.font.color.rgb,
            }
            parts = PLACEHOLDER_RE.split(text)
            run.text = ""
            for part in parts:
                if not part:
                    continue
                nr = p.add_run(part)
                nr.bold = props["bold"]
                nr.italic = props["italic"]
                if props["size"]:
                    nr.font.size = props["size"]
                nr.font.name = props["name"] or "Arial"
                if props["color"]:
                    nr.font.color.rgb = props["color"]
                if PLACEHOLDER_RE.fullmatch(part):
                    nr.font.highlight_color = WD_COLOR_INDEX.YELLOW


def save_doc(doc, filename: str) -> None:
    add_footer(doc)
    highlight_placeholders(doc)
    doc.save(OUT / filename)


def new_company_doc(title: str):
    doc = Document()
    configure_doc(doc, title)
    add_draft_banner(doc)
    add_company_header(doc, title)
    return doc


def build_field_sheet():
    doc = Document()
    configure_doc(doc, "Field Completion and Verification Sheet")
    add_draft_banner(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FIELD COMPLETION AND VERIFICATION SHEET")
    set_run_font(r, size=15, bold=True, color=BLUE)
    add_paragraph(
        doc,
        "Complete this sheet first. Copy the verified values consistently into every "
        "document. Yellow fields must not remain in any document that is finally issued.",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=GREY,
        after=10,
    )

    rows = [
        ("Employee name", EMPLOYEE, "Prefilled; verify spelling against ID"),
        ("Employee ID", "[EMPLOYEE ID]", "Company record"),
        (
            "Employer name for documents",
            COMPANY,
            "Manager must confirm exact registered/trading name on final letterhead",
        ),
        ("Work location", "Aligarh, Uttar Pradesh", "Verify branch"),
        ("Internship period", INTERNSHIP_DATES, "Verified from internship certificate"),
        (
            "Post-internship employment start",
            "[POST-INTERNSHIP START DATE]",
            "Attendance/payroll/manager record; do not assume 17 July 2025",
        ),
        ("Official designation", "[OFFICIAL DESIGNATION]", "Use the recorded title"),
        ("Department/function", "[DEPARTMENT / FUNCTION]", "Company record"),
        ("Employment type", "[FULL-TIME / PART-TIME / TRAINEE / OTHER]", "Use actual status"),
        ("Reporting manager", "[REPORTING MANAGER NAME AND TITLE]", "Manager confirmation"),
        ("Monthly gross remuneration", "₹[GROSS AMOUNT]", "Verify wage/payroll record"),
        ("Monthly deductions", "₹[DEDUCTION AMOUNT OR NIL]", "Verify each month"),
        ("Monthly net cash paid", "₹[NET CASH AMOUNT]", "Verify each month"),
        ("Resignation date", "[RESIGNATION DATE]", "Actual submission date"),
        ("Notice requirement", "[NOTICE PERIOD OR WAIVER]", "Policy/agreement/manager decision"),
        ("Last working day", "[LAST WORKING DATE]", "Written acceptance"),
        ("Full-and-final payment date", "[F&F PAYMENT DATE]", "Insert only after agreed"),
        ("Authorised signatory", "[AUTHORISED SIGNATORY NAME]", "Must independently approve"),
        ("Signatory designation", "[SIGNATORY DESIGNATION]", "Confirm authority"),
        ("BGV contact", "[OFFICIAL EMAIL AND PHONE]", "Obtain consent to list"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Field", "Verified value", "Verification source"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, size=9, color=BLUE)
        shade_cell(table.cell(0, i), LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])
    for label, value, source in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=8.5)
        set_cell_text(cells[1], value, size=8.5)
        set_cell_text(cells[2], source, size=8.2, color=GREY)

    add_paragraph(doc, "RESPONSIBILITIES TO VERIFY", bold=True, size=11, color=BLUE, before=9)
    add_paragraph(
        doc,
        "List only responsibilities personally performed after the internship. Do not copy "
        "internship learning areas into the employment certificate unless the manager confirms "
        "that the same work continued during employment.",
        italic=True,
        size=9.5,
        color=GREY,
    )
    for n in range(1, 7):
        add_bullet(doc, f"[VERIFIED RESPONSIBILITY {n}]")

    add_paragraph(doc, "MANAGER VERIFICATION", bold=True, size=11, color=BLUE, before=8)
    add_paragraph(
        doc,
        "I have checked the above information against available company records and confirm "
        "that the final documents must use only the verified entries.",
        size=9.5,
    )
    add_signatory_block(doc, second=True)
    save_doc(doc, "00_Field_Completion_and_Verification_Sheet.docx")


def build_resignation():
    doc = Document()
    configure_doc(doc, "Resignation Letter")
    add_draft_banner(doc)
    add_paragraph(doc, f"From: {EMPLOYEE}", bold=True, after=2)
    add_paragraph(doc, "Employee ID: [EMPLOYEE ID]", after=2)
    add_paragraph(doc, "Designation: [OFFICIAL DESIGNATION]", after=2)
    add_paragraph(doc, "Department: [DEPARTMENT / FUNCTION]", after=8)
    add_paragraph(doc, "Date: [RESIGNATION DATE]", after=10)
    add_paragraph(doc, "To,", after=2)
    add_paragraph(doc, "[REPORTING MANAGER / AUTHORISED RECIPIENT]", bold=True, after=2)
    add_paragraph(doc, COMPANY, after=2)
    add_paragraph(doc, BRANCH, after=10)
    add_paragraph(
        doc,
        "Subject: Resignation from the position of [OFFICIAL DESIGNATION]",
        bold=True,
        after=10,
    )
    add_paragraph(doc, "Dear [MANAGER NAME],", after=8)
    add_paragraph(
        doc,
        "Please accept this letter as my formal resignation from the position of "
        "[OFFICIAL DESIGNATION] in [DEPARTMENT / FUNCTION]. In accordance with "
        "[APPLICABLE NOTICE PERIOD / AGREED WAIVER], I request that my last working "
        "day be recorded as [PROPOSED LAST WORKING DATE].",
    )
    add_paragraph(
        doc,
        "I will complete a proper handover of my responsibilities, records and company "
        "property. Please confirm my accepted last working day and the designated handover "
        "recipient in writing.",
    )
    add_paragraph(
        doc,
        "I also request the company to issue my resignation acceptance, experience/service "
        "certificate, relieving letter, salary and cash-payment certificate, month-wise wage "
        "statement, full-and-final settlement statement, and no-dues/handover acknowledgement.",
    )
    add_paragraph(
        doc,
        "Thank you for the opportunity and the experience gained during my association with "
        "the organisation.",
    )
    add_paragraph(doc, "Yours sincerely,", before=8, after=1)
    add_signatory_block(doc, employee=True)
    add_paragraph(
        doc,
        "RECEIPT BY COMPANY",
        bold=True,
        size=10,
        color=BLUE,
        before=12,
        after=3,
    )
    add_paragraph(
        doc,
        "Received on [DATE AND TIME] by [NAME AND DESIGNATION]. "
        "Signature/email acknowledgement: [ACKNOWLEDGEMENT DETAILS]",
        size=9,
    )
    save_doc(doc, "01_Resignation_Letter.docx")


def build_acceptance():
    doc = new_company_doc("Resignation Acceptance and Last Working Day Confirmation")
    add_issue_meta(doc)
    add_paragraph(doc, f"To: {EMPLOYEE}", bold=True, after=2)
    add_paragraph(doc, "Employee ID: [EMPLOYEE ID]", after=2)
    add_paragraph(doc, "Designation: [OFFICIAL DESIGNATION]", after=10)
    add_paragraph(doc, f"Dear {EMPLOYEE},", after=8)
    add_paragraph(
        doc,
        "This is to acknowledge receipt of your resignation dated [RESIGNATION DATE] "
        "from the position of [OFFICIAL DESIGNATION] in [DEPARTMENT / FUNCTION].",
    )
    add_paragraph(
        doc,
        "Your resignation has been accepted. After considering [NOTICE PERIOD SERVED / "
        "NOTICE WAIVER / OTHER VERIFIED ARRANGEMENT], your final working day with the "
        f"company will be [LAST WORKING DATE].",
    )
    add_paragraph(
        doc,
        "You are requested to complete the agreed handover, return company property and "
        "submit pending records by [HANDOVER DEADLINE]. Your full-and-final settlement "
        "and separation documents will be processed after completion of the applicable "
        "clearance requirements.",
    )
    add_paragraph(
        doc,
        "Handover recipient: [NAME, DESIGNATION AND CONTACT]",
        bold=True,
        before=6,
    )
    add_signatory_block(doc)
    save_doc(doc, "02_Resignation_Acceptance_and_Last_Working_Day.docx")


def build_employment_confirmation():
    doc = new_company_doc("Post-Internship Employment Status Confirmation")
    add_issue_meta(doc)
    add_paragraph(doc, "TO WHOM IT MAY CONCERN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_paragraph(
        doc,
        f"This letter is issued on [ACTUAL SIGNING DATE] to confirm the records concerning "
        f"{EMPLOYEE}'s association with {COMPANY}.",
    )
    add_paragraph(
        doc,
        f"{EMPLOYEE} completed a Human Resources internship with the company from "
        f"{INTERNSHIP_DATES}. Following the internship, the company records confirm that "
        "he commenced post-internship service on [POST-INTERNSHIP START DATE] as "
        "[OFFICIAL DESIGNATION] in [DEPARTMENT / FUNCTION] on a "
        "[FULL-TIME / PART-TIME / TRAINEE / OTHER] basis.",
    )
    add_paragraph(
        doc,
        "His reporting manager was [REPORTING MANAGER NAME AND TITLE], and his normal "
        "work location was [WORK LOCATION]. His service continued until "
        "[LAST WORKING DATE / PRESENT, IF ISSUED BEFORE EXIT].",
    )
    add_paragraph(
        doc,
        "This is a current-dated confirmation of information verified from available "
        "company records. It is not represented as an appointment letter issued on the "
        "historical effective date.",
        italic=True,
        size=9.5,
        color=GREY,
    )
    add_paragraph(
        doc,
        "For verification, please contact [OFFICIAL COMPANY EMAIL / TELEPHONE].",
        before=8,
    )
    add_signatory_block(doc)
    save_doc(doc, "03_Post_Internship_Employment_Status_Confirmation.docx")


def build_experience():
    doc = new_company_doc("Experience and Service Certificate")
    add_issue_meta(doc)
    add_paragraph(doc, "TO WHOM IT MAY CONCERN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_paragraph(
        doc,
        f"This is to certify that {EMPLOYEE} completed a Human Resources internship with "
        f"{COMPANY} from {INTERNSHIP_DATES}.",
    )
    add_paragraph(
        doc,
        "The company records further confirm that he served as [OFFICIAL DESIGNATION] in "
        "[DEPARTMENT / FUNCTION] from [POST-INTERNSHIP START DATE] to [LAST WORKING DATE] "
        "on a [FULL-TIME / PART-TIME / TRAINEE / OTHER] basis.",
    )
    add_paragraph(doc, "During the post-internship employment period, his verified responsibilities included:")
    for n in range(1, 6):
        add_bullet(doc, f"[VERIFIED EMPLOYMENT RESPONSIBILITY {n}]")
    add_paragraph(
        doc,
        "This certificate records his designation, nature of work and period of service "
        "based on information approved by the authorised signatory.",
    )
    add_paragraph(
        doc,
        "We wish him success in his future endeavours.",
        before=5,
    )
    add_paragraph(
        doc,
        "Employment verification contact: [OFFICIAL EMAIL AND PHONE]",
        size=9.5,
        italic=True,
    )
    add_signatory_block(doc)
    save_doc(doc, "04_Experience_and_Service_Certificate.docx")


def build_relieving():
    doc = new_company_doc("Relieving Letter")
    add_issue_meta(doc)
    add_paragraph(doc, f"To: {EMPLOYEE}", bold=True, after=2)
    add_paragraph(doc, "Employee ID: [EMPLOYEE ID]", after=10)
    add_paragraph(doc, f"Dear {EMPLOYEE},", after=8)
    add_paragraph(
        doc,
        "With reference to your resignation dated [RESIGNATION DATE] and its acceptance, "
        "this is to confirm that you have been relieved from your duties as "
        "[OFFICIAL DESIGNATION] in [DEPARTMENT / FUNCTION] at the close of business on "
        "[LAST WORKING DATE].",
    )
    add_paragraph(
        doc,
        "The company has recorded the completion of the approved handover and return of "
        "company property, subject to the attached/signed no-dues and handover record.",
    )
    add_paragraph(
        doc,
        "Your full-and-final settlement status is [SETTLED ON DATE / PENDING PROCESSING "
        "WITH EXPECTED DATE]. This sentence must accurately match the settlement record.",
    )
    add_paragraph(doc, "We wish you success in your future endeavours.")
    add_signatory_block(doc)
    save_doc(doc, "05_Relieving_Letter.docx")


def build_salary_certificate():
    doc = new_company_doc("Employment, Salary and Cash-Payment Certificate")
    add_issue_meta(doc)
    add_paragraph(doc, "TO WHOM IT MAY CONCERN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_paragraph(
        doc,
        f"This is to certify that {EMPLOYEE}, Employee ID [EMPLOYEE ID], served with "
        f"{COMPANY} as [OFFICIAL DESIGNATION] in [DEPARTMENT / FUNCTION] from "
        "[POST-INTERNSHIP START DATE] to [LAST WORKING DATE].",
    )
    add_paragraph(
        doc,
        "Based on the company's verified wage and attendance records, his regular gross "
        "monthly remuneration was ₹[GROSS MONTHLY AMOUNT], subject to month-specific "
        "attendance, additions and deductions shown in the attached month-wise wage statement.",
    )
    add_paragraph(
        doc,
        "The net remuneration for the verified months listed in that statement was paid "
        "in cash. Each amount and payment date in the statement must be supported by the "
        "company's wage register, cash-payment voucher or equivalent contemporaneous record.",
    )
    add_paragraph(
        doc,
        "This certificate is issued solely as a factual employment and remuneration "
        "confirmation. It must not be issued until the month-wise figures have been checked "
        "against the underlying company records.",
        italic=True,
        size=9.5,
        color=GREY,
    )
    add_paragraph(
        doc,
        "Verification contact: [OFFICIAL COMPANY EMAIL / TELEPHONE]",
        before=8,
    )
    add_signatory_block(doc)
    save_doc(doc, "06_Employment_Salary_and_Cash_Payment_Certificate.docx")


def build_wage_statement():
    doc = new_company_doc("Month-Wise Wage and Cash-Payment Statement")
    add_issue_meta(doc)
    add_paragraph(
        doc,
        f"Employee: {EMPLOYEE}  |  Employee ID: [EMPLOYEE ID]  |  "
        "Designation: [OFFICIAL DESIGNATION]",
        bold=True,
        size=9.5,
        after=7,
    )
    add_paragraph(
        doc,
        "Complete one row for every verified employment month. Delete unused rows. "
        "Figures must be transcribed from attendance, wage and cash-payment records; "
        "do not reconstruct or estimate amounts from memory.",
        italic=True,
        size=8.8,
        color=GREY,
        after=7,
    )
    table = doc.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = [
        "Wage month",
        "Payable days",
        "Gross ₹",
        "Deductions ₹",
        "Net cash ₹",
        "Payment date",
        "Record/voucher ref.",
        "Remarks",
    ]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, size=7.3, color=BLUE)
        shade_cell(table.cell(0, i), LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])
    for n in range(1, 17):
        cells = table.add_row().cells
        values = [
            f"[MONTH {n}]",
            "[DAYS]",
            "[AMOUNT]",
            "[AMOUNT/NIL]",
            "[AMOUNT]",
            "[DATE]",
            "[REFERENCE]",
            "[REMARKS/NIL]",
        ]
        for cell, value in zip(cells, values):
            set_cell_text(cell, value, size=7.2)

    add_paragraph(doc, "TOTALS", bold=True, size=10, color=BLUE, before=8, after=4)
    totals = doc.add_table(rows=3, cols=2)
    totals.style = "Table Grid"
    totals.alignment = WD_TABLE_ALIGNMENT.RIGHT
    for i, (label, value) in enumerate(
        [
            ("Total verified gross remuneration", "₹[TOTAL GROSS]"),
            ("Total verified deductions", "₹[TOTAL DEDUCTIONS]"),
            ("Total verified net cash paid", "₹[TOTAL NET CASH]"),
        ]
    ):
        set_cell_text(totals.cell(i, 0), label, bold=True, size=8.5)
        set_cell_text(totals.cell(i, 1), value, size=8.5)

    add_paragraph(
        doc,
        "Authorised certification: I have verified this statement against available "
        "company attendance, wage and cash-payment records.",
        before=8,
        size=9,
    )
    add_signatory_block(doc, second=True)
    save_doc(doc, "07_Month_Wise_Wage_and_Cash_Payment_Statement.docx")


def build_full_final():
    doc = new_company_doc("Full-and-Final Settlement Statement")
    add_issue_meta(doc)
    add_paragraph(
        doc,
        f"Employee: {EMPLOYEE}  |  Employee ID: [EMPLOYEE ID]  |  "
        "Last working day: [LAST WORKING DATE]",
        bold=True,
        size=9.5,
        after=8,
    )

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Category", "Description/period", "Amount ₹", "Verification/reference"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, size=8.2, color=BLUE)
        shade_cell(table.cell(0, i), LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])
    entries = [
        ("Earning", "Salary through [DATE]", "[AMOUNT]", "[WAGE RECORD]"),
        ("Earning", "Pending prior salary, if any", "[AMOUNT/NIL]", "[REFERENCE]"),
        ("Earning", "Approved reimbursement", "[AMOUNT/NIL]", "[REFERENCE]"),
        ("Earning", "Leave encashment, if applicable", "[AMOUNT/NIL]", "[REFERENCE]"),
        ("Earning", "Incentive/other approved earning", "[AMOUNT/NIL]", "[REFERENCE]"),
        ("Deduction", "Statutory deduction, if applicable", "[AMOUNT/NIL]", "[REFERENCE]"),
        ("Deduction", "Advance/loan recovery", "[AMOUNT/NIL]", "[REFERENCE]"),
        ("Deduction", "Notice-period adjustment", "[AMOUNT/NIL]", "[REFERENCE]"),
        ("Deduction", "Other authorised deduction", "[AMOUNT/NIL]", "[REFERENCE]"),
    ]
    for row in entries:
        cells = table.add_row().cells
        for c, value in zip(cells, row):
            set_cell_text(c, value, size=8)

    add_paragraph(doc, "SETTLEMENT SUMMARY", bold=True, size=10, color=BLUE, before=8, after=4)
    summary = doc.add_table(rows=5, cols=2)
    summary.style = "Table Grid"
    summary.alignment = WD_TABLE_ALIGNMENT.RIGHT
    summary_rows = [
        ("Total earnings", "₹[TOTAL EARNINGS]"),
        ("Total deductions", "₹[TOTAL DEDUCTIONS]"),
        ("Net amount payable", "₹[NET PAYABLE]"),
        ("Actual payment date and mode", "[DATE AND CASH/BANK/OTHER]"),
        ("Payment/voucher reference", "[REFERENCE]"),
    ]
    for i, (label, value) in enumerate(summary_rows):
        set_cell_text(summary.cell(i, 0), label, bold=True, size=8.5)
        set_cell_text(summary.cell(i, 1), value, size=8.5)

    add_paragraph(
        doc,
        "Employee acknowledgement — sign only after receipt:",
        bold=True,
        size=9.5,
        before=9,
        after=4,
    )
    add_paragraph(
        doc,
        f"I, {EMPLOYEE}, acknowledge receipt of ₹[ACTUAL AMOUNT RECEIVED] on "
        "[ACTUAL RECEIPT DATE] by [PAYMENT MODE]. This acknowledgement relates only to "
        "the itemised amounts above and must not be signed before payment is received.",
        size=9,
    )
    add_signatory_block(doc, second=True)
    save_doc(doc, "08_Full_and_Final_Settlement_Statement.docx")


def build_handover():
    doc = new_company_doc("Handover and No-Dues Record")
    add_issue_meta(doc)
    add_paragraph(
        doc,
        f"Employee: {EMPLOYEE}  |  Employee ID: [EMPLOYEE ID]  |  "
        "Designation: [OFFICIAL DESIGNATION]",
        bold=True,
        size=9.5,
        after=8,
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Item/area", "Details", "Returned/handed over", "Recipient", "Date/signature"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, size=8, color=BLUE)
        shade_cell(table.cell(0, i), LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])
    items = [
        "Pending work and status",
        "HR files/records",
        "Passwords/access transferred through approved process",
        "Employee/candidate matters requiring follow-up",
        "Laptop/desktop/accessories",
        "Employee ID/access card/keys",
        "Documents/registers/stationery",
        "Advances/reimbursements",
        "Other company property",
    ]
    for item in items:
        cells = table.add_row().cells
        values = [item, "[DETAILS/NIL]", "[YES/NO/N.A.]", "[NAME]", "[DATE/SIGN]"]
        for c, value in zip(cells, values):
            set_cell_text(c, value, size=7.8)

    add_paragraph(
        doc,
        "CLEARANCE OUTCOME",
        bold=True,
        size=10,
        color=BLUE,
        before=8,
        after=4,
    )
    add_paragraph(
        doc,
        "Based on the itemised record above, the employee's handover status is "
        "[COMPLETE / COMPLETE WITH LISTED EXCEPTIONS / PENDING]. Any exception must be "
        "described here: [DETAILS OR NIL].",
        size=9,
    )
    add_paragraph(
        doc,
        "This record confirms only the specific handover and property items listed above. "
        "It is not a blanket waiver of unrecorded salary or statutory entitlements.",
        italic=True,
        size=8.8,
        color=GREY,
    )
    add_signatory_block(doc, second=True)
    save_doc(doc, "09_Handover_and_No_Dues_Record.docx")


def build_recommendation():
    doc = new_company_doc("Manager Recommendation and Professional Reference")
    add_issue_meta(doc)
    add_paragraph(doc, "TO WHOM IT MAY CONCERN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_paragraph(
        doc,
        f"I am pleased to provide this professional reference for {EMPLOYEE}, who served "
        f"with {COMPANY} as [OFFICIAL DESIGNATION] in [DEPARTMENT / FUNCTION] from "
        "[POST-INTERNSHIP START DATE] to [LAST WORKING DATE]. I was his "
        "[REPORTING RELATIONSHIP] during this period.",
    )
    add_paragraph(
        doc,
        "His verified responsibilities included [TWO OR THREE SPECIFIC RESPONSIBILITIES]. "
        "A representative example of his contribution was [SPECIFIC, FACTUAL EXAMPLE "
        "WITHOUT CONFIDENTIAL INFORMATION].",
    )
    add_paragraph(
        doc,
        "Based on my direct observation, his demonstrated strengths included "
        "[TWO OR THREE OBSERVED STRENGTHS]. He handled [RELEVANT TYPE OF WORK] with "
        "[ACCURATE DESCRIPTION OF QUALITY / RELIABILITY / PROFESSIONALISM].",
    )
    add_paragraph(
        doc,
        "I would consider him for entry-level opportunities in [RECOMMENDED FUNCTIONAL "
        "DOMAIN] where these capabilities are relevant.",
    )
    add_paragraph(
        doc,
        "I consent to being contacted for factual employment verification at "
        "[OFFICIAL EMAIL / TELEPHONE].",
        before=6,
    )
    add_signatory_block(doc)
    save_doc(doc, "10_Manager_Recommendation_and_Reference.docx")


def build_receipt():
    doc = new_company_doc("Separation Documents Issuance and Receipt Record")
    add_issue_meta(doc)
    add_paragraph(
        doc,
        f"Employee: {EMPLOYEE}  |  Employee ID: [EMPLOYEE ID]  |  "
        "Last working day: [LAST WORKING DATE]",
        bold=True,
        size=9.5,
        after=8,
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Document", "Issued?", "Issue date", "Original/PDF", "Remarks/reference"]
    for i, h in enumerate(headers):
        set_cell_text(table.cell(0, i), h, bold=True, size=8, color=BLUE)
        shade_cell(table.cell(0, i), LIGHT_BLUE)
    set_repeat_table_header(table.rows[0])
    docs = [
        "Resignation acceptance",
        "Post-internship employment status confirmation",
        "Experience/service certificate",
        "Relieving letter",
        "Employment, salary and cash-payment certificate",
        "Month-wise wage and cash-payment statement",
        "Full-and-final settlement statement",
        "Handover and no-dues record",
        "Manager recommendation/reference",
        "PF/UAN/ESI information, if applicable",
        "Other: [DOCUMENT]",
    ]
    for item in docs:
        cells = table.add_row().cells
        values = [item, "[YES/NO/N.A.]", "[DATE]", "[FORMAT]", "[REMARKS]"]
        for c, value in zip(cells, values):
            set_cell_text(c, value, size=7.8)
    add_paragraph(
        doc,
        "The employee acknowledges receipt only of documents marked “Yes” above. This "
        "receipt does not confirm payment of any amount unless separately recorded in the "
        "full-and-final settlement.",
        italic=True,
        size=8.8,
        color=GREY,
        before=8,
    )
    add_signatory_block(doc, second=True)
    save_doc(doc, "11_Separation_Documents_Issuance_and_Receipt.docx")


def build_readme():
    content = f"""# MyTVS employment exit document pack

Prepared for: **{EMPLOYEE}**  
Employer name requested for these drafts: **{COMPANY}**  
Known internship period: **{INTERNSHIP_DATES}**

## Important status

Every Word file is an editable **draft for independent manager/authorised-signatory review**.
Yellow fields are unverified. No company-issued document is valid until the facts are checked
against company records and an authorised person signs it on approved letterhead.

This pack does not create or simulate signatures, stamps, salary vouchers, historical payroll
records or a backdated appointment letter. The post-internship confirmation is intentionally
current-dated: it confirms an earlier effective date only after the manager checks the company's
attendance, wage or other contemporaneous records.

Because Mohammad has been acting in an HR capacity, his own separation documents should be
independently reviewed and approved. He should not be both the sole preparer/approver and the
purported company issuer.

## Recommended order

1. Complete `00_Field_Completion_and_Verification_Sheet.docx`.
2. Confirm the exact registered/trading employer name used on company letterhead, designation,
   employment type, start date, cash salary and last working day.
3. Submit `01_Resignation_Letter.docx` and obtain a dated acknowledgement.
4. Have the manager issue the acceptance and confirm the last working day.
5. Populate the wage statement only from attendance, wage-register and cash-payment records.
6. Complete and sign the handover record.
7. Calculate the full-and-final settlement and sign the employee acknowledgement only after
   the stated money is actually received.
8. Issue the experience, relieving, salary and recommendation documents with the actual issue
   date, approved letterhead, authorised signature and company stamp where normally used.
9. Use the document-receipt form to record what was actually provided.

## Files

- `00_Field_Completion_and_Verification_Sheet.docx`
- `01_Resignation_Letter.docx`
- `02_Resignation_Acceptance_and_Last_Working_Day.docx`
- `03_Post_Internship_Employment_Status_Confirmation.docx`
- `04_Experience_and_Service_Certificate.docx`
- `05_Relieving_Letter.docx`
- `06_Employment_Salary_and_Cash_Payment_Certificate.docx`
- `07_Month_Wise_Wage_and_Cash_Payment_Statement.docx`
- `08_Full_and_Final_Settlement_Statement.docx`
- `09_Handover_and_No_Dues_Record.docx`
- `10_Manager_Recommendation_and_Reference.docx`
- `11_Separation_Documents_Issuance_and_Receipt.docx`

## Facts still required

- Exact registered/trading employer name used on approved letterhead
- Official employee ID
- Exact post-internship employment start date
- Official designation and department
- Employment type
- Reporting manager and authorised signatory
- Actual month-wise gross salary, deductions, net cash paid and payment dates
- Notice requirement, resignation date and accepted last working day
- Verified post-internship responsibilities and achievements
- Official company contact that will answer background-verification requests

## Recordkeeping cautions

- Do not change the internship dates already supported by the certificate.
- Do not estimate salary months or create retrospective wage slips without source records.
- Do not describe the whole period as an internship if the company verifies a separate
  post-internship role.
- Do not use a preferred résumé title in place of the official designation. A functional
  clarification can be added separately later if it is accurate.
- Do not sign blank forms or acknowledge full-and-final payment before receiving it.
- Retain colour scans of every signed document and the email transmitting each document.
"""
    (OUT / "README.md").write_text(content, encoding="utf-8")


def build_completion_checklist():
    content = """document,status,critical_checks
00 Field Completion and Verification Sheet,DRAFT,"All yellow fields completed; sources checked; manager signs verification"
01 Resignation Letter,DRAFT,"Actual submission date; proposed last day; company receipt/email acknowledgement"
02 Resignation Acceptance,DRAFT,"Accepted last day; notice decision; authorised signature"
03 Employment Status Confirmation,DRAFT,"Current issue date; verified historical start date; no false backdating"
04 Experience and Service Certificate,DRAFT,"Separate internship/employment; official title; true responsibilities"
05 Relieving Letter,DRAFT,"Issued only after actual release; settlement wording accurate"
06 Salary and Cash Payment Certificate,DRAFT,"Amounts match wage records; cash mode explicitly confirmed"
07 Month-Wise Wage Statement,DRAFT,"Every month supported; totals checked; unused rows deleted"
08 Full-and-Final Settlement,DRAFT,"Itemised; deductions authorised; employee signs only after receipt"
09 Handover and No-Dues,DRAFT,"All items specific; exceptions recorded; no blanket wage waiver"
10 Manager Recommendation,DRAFT,"Direct observations only; no invented metrics; manager consents to contact"
11 Document Receipt,DRAFT,"Only actually issued documents marked Yes"
"""
    (OUT / "DOCUMENT_COMPLETION_CHECKLIST.csv").write_text(content, encoding="utf-8")


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    build_field_sheet()
    build_resignation()
    build_acceptance()
    build_employment_confirmation()
    build_experience()
    build_relieving()
    build_salary_certificate()
    build_wage_statement()
    build_full_final()
    build_handover()
    build_recommendation()
    build_receipt()
    build_readme()
    build_completion_checklist()
    print(f"Generated exit pack in: {OUT}")


if __name__ == "__main__":
    main()
