#!/usr/bin/env python3
"""Build the 5 August 2026 nine-role remote application CV pack."""

from __future__ import annotations

import copy
import csv
import importlib.util
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
BASE_SOURCE = ROOT / "98_Maintenance" / "generate_remote_targeted_resumes.py"
OLD_SOURCE = ROOT / "98_Maintenance" / "generate_high_paying_early_career_resumes.py"
PDF_RENDERER = ROOT / "98_Maintenance" / "render_remote_targeted_resumes_pdf.py"
OUT = ROOT / "07_Remote_Job_Applications" / "2026-08-05_Current_High_Paying_Remote"


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


BASE = load_module(BASE_SOURCE, "base_remote_resumes_aug05")
OLD = load_module(OLD_SOURCE, "old_high_paying_resumes_aug05")
RENDERER = load_module(PDF_RENDERER, "pdf_renderer_aug05")

CONTACT = BASE.CONTACT
EDUCATION = BASE.EDUCATION
LEARNING = BASE.LEARNING
AI_TOOLS = BASE.AI_TOOLS
SURE = copy.deepcopy(BASE.SURE)
MYTVS_HR_OPS = copy.deepcopy(BASE.MYTVS_HR_OPS)


TVS_HR_INTERNSHIP = {
    "title": "Human Resources Intern",
    "org": "TVS Automobile Solutions Pvt. Ltd. · Aligarh, India",
    "dates": "26 May 2025 – 16 July 2025",
    "bullets": [
        "Completed a seven-week HR internship with verified exposure to recruitment and selection, employee "
        "engagement, performance evaluation, and HR documentation.",
    ],
}

SURE_COMPACT = {
    "title": "Consultant, SURE Program",
    "org": "Aligarh Muslim University × C. T. Bauer College of Business, University of Houston",
    "dates": "August 2024 – October 2024",
    "bullets": [
        "Completed a 12-week consultant programme and learned a structured business-plan preparation process "
        "in an AMU–University of Houston collaboration.",
    ],
}

AI_WORKFLOW_PROJECT = {
    "title": "Evidence-Controlled Career Application System",
    "org": "Independent AI-assisted project · Claude, ChatGPT Codex, Visual Studio Code",
    "dates": "2026",
    "bullets": [
        "Built an auditable workspace that maps source records to approved claims, creates nine role-specific "
        "application packs, and runs content and file-quality checks before use.",
    ],
}


def mytvs_experience(bullets: list[str]) -> list[dict]:
    """Return a role-targeted but evidence-controlled experience sequence."""
    current = {
        "title": "HR Operations",
        "org": "myTVS · Aligarh, India",
        "dates": "July 2025 – Present",
        "bullets": bullets,
    }
    return [current, copy.deepcopy(TVS_HR_INTERNSHIP), copy.deepcopy(SURE_COMPACT)]


def old_job(company: str, role: str) -> dict:
    for job in OLD.JOBS:
        if job["company"] == company and job["role"] == role:
            return copy.deepcopy(job)
    raise KeyError(f"Old job not found: {company} — {role}")


def configure_old(company: str, role: str, order: int, slug: str, status: str) -> dict:
    job = old_job(company, role)
    job.update(order=order, slug=slug, category="priority", status=status)
    return job


EVERIS = configure_old(
    "Everis",
    "People Operations Specialist",
    1,
    "01_Everis_People_Operations_Specialist",
    "Apply first — India expressly eligible; salary disclosed",
)
EVERIS.update(
    special_instruction=(
        "Use the Ashby-targeted PDF. The application is output-focused and competitive; do not add "
        "HRIS, payroll, onboarding ownership, ATS, or quantified results unless evidence is supplied."
    ),
    supported_keywords=[
        "People Operations", "HR administration", "HR documentation", "employee communication",
        "management follow-up", "confidential information", "AI tools", "remote operations",
        "organisation", "independent prioritisation", "follow-through",
    ],
    gap_keywords=[
        "HRIS administration", "payroll ownership", "leave administration", "onboarding ownership",
        "offboarding ownership", "ATS administration", "SOP ownership",
    ],
    skills=[
        "Day-to-day People Operations and HR administration",
        "Employee communication and management follow-up",
        "HR documentation and confidential information handling",
        "Recruitment and selection exposure",
        "Structured documentation and information organisation",
        "Microsoft Excel, Word, and PowerPoint",
        "AI tools for research, drafting, and output review",
        "Remote operations readiness, independent prioritisation, and follow-through",
    ],
)


QUANTUM = {
    "order": 2,
    "slug": "02_QuantumLoopAI_Junior_Admin_Operations_Associate",
    "category": "priority",
    "company": "QuantumLoopAI",
    "role": "Junior Admin Operations Associate — India Remote",
    "url": "https://apply.workable.com/quantumloopai/j/251EFE0C83/apply/",
    "platform": "Workable",
    "fit": "Excellent entry-level operations fit",
    "status": "Apply immediately — India remote; no prior experience required",
    "headline": "ADMIN OPERATIONS | PROCESS ACCURACY | DOCUMENTATION & TRACKERS",
    "summary": (
        "Early-career operations professional with approximately one year of post-internship HR Operations "
        "experience in a lean organisation of about 70 employees. Brings structured administrative support, "
        "careful documentation, organised employee information, routine coordination, management follow-up, "
        "confidentiality, and independent prioritisation. Comfortable with Microsoft Office and hands-on use "
        "of modern AI tools for research, drafting, information organisation, and output review."
    ),
    "skills": [
        "Administrative and day-to-day operations support",
        "Structured processes and rule-based workflows",
        "Documentation, records, and tracker organisation",
        "Attention to detail, accuracy checks, and inconsistency awareness",
        "Cross-team communication and management follow-up",
        "Microsoft Excel, Word, and PowerPoint",
        "AI, automation, and technology-product interest",
        "Remote-work discipline and independent learning",
    ],
    "experience": [MYTVS_HR_OPS, SURE],
    "learning": ["excel", "performance", "uci", "powerbi"],
    "tools": AI_TOOLS,
    "company_analysis": (
        "QuantumLoopAI seeks an India-based beginner for structured operational work behind an AI healthtech "
        "platform. The description repeatedly prioritises rule-based processes, accurate records and trackers, "
        "UK English, coordination, error spotting, reliability, everyday digital tools, AI curiosity, and UK "
        "business-hour availability; it explicitly says previous experience is not essential."
    ),
    "supported_keywords": [
        "admin operations", "day-to-day operations", "structured processes", "rule-based workflows",
        "documentation", "records", "trackers", "accuracy", "attention to detail", "coordination",
        "follow-up", "Excel", "AI", "remote work",
    ],
    "gap_keywords": [
        "CRM administration", "project-management platforms", "structured dashboards",
        "professional automation", "healthtech operations",
    ],
    "keywords": [
        "admin operations", "structured processes", "rule-based workflows", "documentation", "trackers",
        "accuracy", "coordination", "follow-up", "Excel", "AI", "remote work",
    ],
    "gaps": [
        "No verified production CRM, project-management platform, or dashboard experience.",
        "No verified healthtech or UK-company experience.",
        "The application separately requires confirmation of UK working hours, equipment, internet, salary, "
        "and candidate-authored process/error examples.",
    ],
    "pitch": (
        "My lean-company HR Operations experience transfers directly to structured administration, accurate "
        "records, recurring follow-up, confidentiality, and dependable process execution. I also have genuine "
        "hands-on familiarity with AI tools and can learn QuantumLoopAI's systems without claiming prior "
        "healthtech or CRM experience."
    ),
    "special_instruction": (
        "The Workable form asks for a detailed process example and an error-detection example. Write those in "
        "your own words and make sure the salary, equipment, internet, and UK-hours answers are accurate."
    ),
}


OUTCAPPED = configure_old(
    "Outcapped",
    "Operations Associate",
    3,
    "03_Outcapped_Operations_Associate",
    "Priority stretch — one-year threshold met; worldwide remote",
)
OUTCAPPED.update(
    supported_keywords=[
        "operations associate", "administration", "documentation", "information organisation",
        "quality review", "accuracy", "follow-through", "process improvement", "AI tools",
        "remote operations", "attention to detail", "owner thinking",
    ],
    gap_keywords=[
        "Notion", "Airtable", "Google Workspace", "Slack", "Zapier", "Make",
        "light bookkeeping", "invoice processing", "reconciliations", "website upkeep",
    ],
    skills=[
        "Administrative and day-to-day operations support",
        "Documentation and information organisation",
        "Quality review, accuracy, and attention to detail",
        "Management follow-up and action completion",
        "Process-improvement mindset, owner thinking, and reliable follow-through",
        "Microsoft Excel, Word, and PowerPoint",
        "AI tools and remote operations readiness",
    ],
)


RECRUIT_CRM = {
    "order": 4,
    "slug": "04_Recruit_CRM_Customer_Success_Associate_2026",
    "category": "priority",
    "company": "Recruit CRM",
    "role": "Customer Success Associate — 2026 India",
    "url": "https://careers.recruitcrm.io/17712190853810129269dFC",
    "platform": "Recruit CRM careers ATS",
    "fit": "Strong 2026-graduate customer-success transition",
    "status": "Apply now — exclusively 2026 graduates/postgraduates; remote India",
    "headline": "CUSTOMER SUCCESS | PROFESSIONAL COMMUNICATION | SaaS LEARNING",
    "summary": (
        "MHRM candidate completing in 2026, Economics graduate, and early-career HR Operations professional "
        "with approximately one year of experience in a lean organisation of about 70 employees. Brings "
        "professional employee communication, clear follow-up, organised documentation, problem-solving, "
        "confidentiality, and independent learning. Interested in transferring these strengths to SaaS "
        "customer success and helping international users understand and adopt Recruit CRM."
    ),
    "skills": [
        "Professional communication and relationship awareness",
        "Employee queries, follow-up, and clear information sharing",
        "Documentation and information organisation",
        "Problem-solving and critical thinking",
        "Software learning and SaaS product curiosity",
        "Microsoft Excel, Word, and PowerPoint",
        "AI-assisted research, drafting, and output review",
        "Customer-service interest, service delivery, consistency, and integrity",
    ],
    "experience": [MYTVS_HR_OPS, SURE],
    "learning": ["performance", "excel", "uci", "powerbi"],
    "tools": AI_TOOLS,
    "company_analysis": (
        "Recruit CRM's current 2026 India programme is a six-month customer-success traineeship leading to a "
        "possible full-time associate offer. The role involves learning the SaaS platform, answering client "
        "queries through chat, email, and video, explaining features and workarounds, communicating customer "
        "needs, building relationships, and delivering consistent service. Selection tests English, aptitude, "
        "analytical thinking, communication, and problem-solving."
    ),
    "supported_keywords": [
        "customer success", "professional communication", "follow-up", "documentation", "problem-solving",
        "software learning", "SaaS", "relationship awareness", "service", "consistency", "integrity",
        "2026 postgraduate",
    ],
    "gap_keywords": [
        "CRM administration", "client portfolio", "chat support", "email support", "video-call support",
        "feature troubleshooting", "customer onboarding", "renewals", "customer retention",
    ],
    "keywords": [
        "customer success", "SaaS", "communication", "follow-up", "problem-solving", "documentation",
        "software learning", "service delivery", "consistency", "integrity",
    ],
    "gaps": [
        "No verified external SaaS customer-support, chat, email-ticket, video-support, or CRM experience.",
        "The first six months are a traineeship; the full-time offer and advertised CTC are performance-dependent.",
        "The candidate must be completing graduation or postgraduation in 2026 and must confirm this accurately.",
    ],
    "pitch": (
        "My employee-facing HR Operations work has developed communication, patience, follow-up, documentation, "
        "integrity, and independent problem-solving. My MHRM completion year matches the 2026 eligibility rule, "
        "and my technology curiosity supports a truthful transition into SaaS customer success."
    ),
    "special_instruction": (
        "Do not describe internal employee communication as external customer-support experience. Prepare for "
        "the English/aptitude test and confirm willingness for the initial 12:00–22:00 IST training shift."
    ),
}


CANONICAL_HR = configure_old(
    "Canonical",
    "Graduate HR Generalist – APAC",
    5,
    "05_Canonical_Graduate_HR_Generalist_APAC",
    "Priority high-competition application — APAC remote",
)
CANONICAL_HR.update(
    supported_keywords=[
        "graduate HR", "HR operations", "employee support", "HR documentation", "employee information",
        "confidential data", "accuracy", "communication", "prioritisation", "adaptability", "integrity",
        "process improvement", "independent work",
    ],
    gap_keywords=[
        "HRIS administration", "global benefits", "payroll inputs", "purchase orders", "GDPR audits",
        "global projects", "country-specific onboarding",
    ],
    skills=[
        "Day-to-day HR operations and employee support",
        "HR documentation and organised employee information",
        "Employee communication and management follow-up",
        "Recruitment, engagement, and performance-management exposure",
        "Integrity, confidential data handling, and accuracy",
        "Microsoft Excel, Word, and PowerPoint",
        "Process-improvement awareness, independent work, prioritisation, and adaptability",
    ],
)


MERIDIAL_BUSINESS = configure_old(
    "Meridial",
    "Business and Management Specialist – Freelance AI Trainer",
    6,
    "06_Meridial_Business_Management_AI_Trainer",
    "Apply — worldwide entry-level freelance project; rate and hours variable",
)
MERIDIAL_BUSINESS.update(
    supported_keywords=[
        "business management", "management principles", "organisational behaviour", "business operations",
        "performance management", "economics", "business reasoning", "decision-making", "critical reading",
        "structured explanation", "AI output review", "logical consistency",
    ],
    gap_keywords=[
        "professional AI training", "model evaluation employment", "prompt evaluation framework ownership",
        "published business research", "strategy consulting",
    ],
    special_instruction=(
        "Greenhouse accepts PDF, DOCX, DOC, RTF, and TXT. Use the PDF, select only genuinely proficient "
        "languages, and do not treat the advertised maximum hourly rate as an expected offer."
    ),
    skills=[
        "Business management and management principles",
        "Organisational behaviour and performance-management coursework",
        "Business operations context from lean-company HR Operations",
        "Economics, decision-making, and business reasoning foundation",
        "Critical reading, logical consistency, and structured explanation",
        "AI output review using Claude and ChatGPT Codex",
        "Microsoft Excel, Word, and PowerPoint",
    ],
)


CANONICAL_CSM = configure_old(
    "Canonical",
    "Graduate Customer Success Manager",
    7,
    "07_Canonical_Graduate_Customer_Success_Manager",
    "High-value stretch — worldwide remote graduate role",
)
CANONICAL_CSM.update(
    supported_keywords=[
        "customer success", "empathetic communication", "follow-up", "documentation", "organisation",
        "problem-solving", "critical thinking", "independent learning", "technology interest",
        "cross-functional teamwork",
    ],
    gap_keywords=[
        "customer portfolio", "support tickets", "SaaS onboarding", "product adoption", "retention",
        "churn risk", "Ubuntu", "Linux", "cloud infrastructure",
    ],
    skills=[
        "Employee communication and issue follow-up",
        "Empathetic communication and relationship awareness",
        "Documentation and information organisation",
        "Problem-solving and critical thinking",
        "Microsoft Excel, Word, and PowerPoint",
        "AI-assisted research, drafting, and output review",
        "Technology interest, independent learning, and cross-functional teamwork",
    ],
)


CANONICAL_SDR = configure_old(
    "Canonical",
    "Graduate Sales Development Representative",
    8,
    "08_Canonical_Graduate_Sales_Development_Representative",
    "High-value stretch — worldwide remote graduate role",
)
CANONICAL_SDR.update(
    supported_keywords=[
        "sales development", "business development", "business research", "market research",
        "professional communication", "follow-up", "economics", "business reasoning", "technology interest",
        "curiosity", "persistence", "continuous learning",
    ],
    gap_keywords=[
        "cold calling", "prospecting", "CRM", "lead generation", "quota", "pipeline ownership",
        "outbound campaigns", "inbound sales", "Ubuntu", "Linux", "cloud technology",
    ],
    skills=[
        "Business research, market research, and information synthesis",
        "Professional communication and structured follow-up",
        "Economics foundation and business reasoning",
        "Business-plan preparation process",
        "Microsoft Excel, Word, and PowerPoint",
        "AI-assisted research, drafting, and output review",
        "Technology interest, curiosity, persistence, and continuous learning",
    ],
)


MERIDIAL_INVESTMENT = configure_old(
    "Meridial",
    "Investment Specialist – Freelance AI Trainer",
    9,
    "09_Meridial_Investment_AI_Trainer",
    "Selective stretch — worldwide entry-level freelance project",
)
MERIDIAL_INVESTMENT.update(
    supported_keywords=[
        "economics", "investment fundamentals", "financial reasoning", "macroeconomics", "risk concepts",
        "critical reading", "structured explanation", "AI output review", "logical consistency",
        "analytical thinking",
    ],
    gap_keywords=[
        "asset management", "brokerage", "financial modelling", "valuation modelling", "portfolio management",
        "derivatives", "professional investment research", "professional AI training",
    ],
    special_instruction=(
        "Greenhouse accepts PDF, DOCX, DOC, RTF, and TXT. Apply only after revising investments, valuation, "
        "portfolio construction, markets, and risk; do not imply professional finance experience."
    ),
    skills=[
        "Economics, macroeconomics, and market-reasoning foundation",
        "Interest in investment fundamentals, financial reasoning, and risk concepts",
        "Performance measurement and analytics coursework",
        "Business-plan preparation process",
        "Critical reading, logical consistency, and structured explanation",
        "AI output review using Claude and ChatGPT Codex",
        "Careful interpretation of evidence and limitations",
    ],
)


JOBS = [
    EVERIS,
    QUANTUM,
    OUTCAPPED,
    RECRUIT_CRM,
    CANONICAL_HR,
    MERIDIAL_BUSINESS,
    CANONICAL_CSM,
    CANONICAL_SDR,
    MERIDIAL_INVESTMENT,
]


# Senior-recruiter rewrite: every version leads with role evidence, uses only supported
# metrics, and carries role-specific achievements into the experience section itself.
RECRUITER_REWRITES = {
    1: {
        "headline": "PEOPLE OPERATIONS | ~70-EMPLOYEE WORKFORCE | AI-ASSISTED SYSTEMS",
        "summary": (
            "People Operations professional with approximately one year supporting routine HR administration "
            "for a workforce of about 70 employees. Coordinates employee communication, confidential "
            "documentation, management follow-ups, and recurring priorities independently in a lean "
            "environment. Built an AI-assisted, evidence-controlled documentation workflow; currently pursuing "
            "an MHRM and scored 90.4% in Performance Management."
        ),
        "skills": [
            "People Operations", "HR administration", "employee communication",
            "documentation and templates", "confidential information", "management follow-up",
            "independent prioritisation", "follow-through", "remote operations readiness",
        ],
        "experience": mytvs_experience([
            "Coordinate recurring People Operations support for about 70 employees, handling routine HR "
            "administration, employee communication, and management follow-ups from request to completion.",
            "Organise employee documents and information for accurate, confidential, and timely retrieval.",
            "Prioritise changing HR requirements independently in a lean environment and keep outstanding "
            "actions visible until completion.",
        ]),
        "learning": ["performance", "excel", "uci", "vitara"],
        "technical": (
            "Microsoft Excel, Word, and PowerPoint · AI tools: Claude and ChatGPT Codex for research, drafting, "
            "and output review · Visual Studio Code for structured information organisation"
        ),
        "projects": [copy.deepcopy(AI_WORKFLOW_PROJECT)],
    },
    2: {
        "headline": "ADMIN OPERATIONS | PROCESS ACCURACY | AI & DIGITAL TOOLS",
        "summary": (
            "Early-career operations professional with approximately one year executing structured HR "
            "administration for about 70 employees. Maintains organised, confidential records; coordinates "
            "updates and follow-ups; and checks actions through completion. Comfortable with Excel, Word, AI "
            "tools, and independent digital work, with strong process discipline and interest in AI-enabled "
            "systems."
        ),
        "skills": [
            "admin operations", "day-to-day operations", "structured processes", "rule-based workflows",
            "documentation, records, and trackers", "accuracy review", "attention to detail",
            "coordination and follow-up",
        ],
        "experience": mytvs_experience([
            "Execute recurring HR administration and employee coordination for a workforce of about 70, "
            "following structured routines and management follow-ups through to completion.",
            "Maintain confidential employee documentation and organised records with attention to accuracy, "
            "completeness, and retrieval readiness.",
            "Prioritise recurring tasks independently, flag pending actions, and adapt to changing operational "
            "requirements in a lean environment.",
        ]),
        "learning": ["excel", "powerbi", "performance"],
        "technical": (
            "Microsoft Excel, Word, and PowerPoint · Claude, ChatGPT Codex, and Visual Studio Code for research, "
            "drafting, information organisation, and output review · introductory Power BI workshop"
        ),
        "projects": [copy.deepcopy(AI_WORKFLOW_PROJECT)],
    },
    3: {
        "headline": "OPERATIONS | DOCUMENTATION | QUALITY & FOLLOW-THROUGH",
        "summary": (
            "Operations professional with approximately one year supporting recurring administration for a "
            "lean, 70-person organisation. Brings structured documentation, information organisation, quality "
            "review, management follow-up, confidentiality, and independent ownership of priorities. Uses AI "
            "tools and Visual Studio Code to build clear, repeatable, evidence-controlled workflows."
        ),
        "skills": [
            "operations support", "administration", "documentation", "information organisation",
            "quality assurance and accuracy", "attention to detail", "process improvement", "owner thinking",
            "remote operations readiness", "reliable follow-through",
        ],
        "experience": mytvs_experience([
            "Keep recurring HR administration, employee communication, and management follow-ups moving for a "
            "workforce of about 70 employees.",
            "Maintain confidential documents and organise employee information for accurate, timely retrieval "
            "and quality review.",
            "Take ownership of changing priorities in a lean environment and follow outstanding actions through "
            "to completion.",
        ]),
        "learning": ["excel", "performance", "powerbi"],
        "technical": (
            "Microsoft Excel, Word, and PowerPoint · Claude, ChatGPT Codex, and Visual Studio Code for research, "
            "drafting, information organisation, and output review · introductory Power BI workshop"
        ),
        "projects": [copy.deepcopy(AI_WORKFLOW_PROJECT)],
    },
    4: {
        "headline": "CUSTOMER SUCCESS | EMPLOYEE COMMUNICATION | SaaS LEARNING",
        "summary": (
            "2026 postgraduate candidate completing an MHRM, with approximately one year of HR Operations experience "
            "supporting about 70 employees through clear communication, follow-up, and organised documentation. "
            "Brings patience, problem-solving, software curiosity, and hands-on AI-tool use. Ready to learn "
            "Recruit CRM and deliver clear, consistent support to international users."
        ),
        "skills": [
            "professional communication", "employee support", "issue follow-up", "problem-solving",
            "documentation", "relationship awareness", "service consistency", "integrity", "software learning",
        ],
        "experience": mytvs_experience([
            "Support routine employee communication and HR administration for about 70 employees, coordinating "
            "management follow-ups and keeping actions moving to completion.",
            "Organise confidential employee documentation so information remains accurate and easy to retrieve.",
            "Work independently, prioritise recurring requests, and communicate updates clearly through "
            "completion.",
        ]),
        "learning": ["performance", "excel", "uci"],
        "technical": (
            "Microsoft Excel, Word, and PowerPoint · Claude, ChatGPT Codex, and Visual Studio Code for research, "
            "drafting, information organisation, and output review"
        ),
        "projects": [],
    },
    5: {
        "headline": "HR OPERATIONS | ~70-EMPLOYEE WORKFORCE | ACCURACY & INTEGRITY",
        "summary": (
            "Graduate HR professional with approximately one year supporting routine people administration for "
            "about 70 employees. Brings employee communication, confidential records handling, management "
            "follow-up, independent prioritisation, accuracy, and adaptability. Pursuing an MHRM; scored 90.4% "
            "in Performance Management and 78.8% in HR Analytics Using Excel."
        ),
        "skills": [
            "HR operations", "employee support", "HR documentation", "employee information",
            "confidential data", "accuracy", "workload prioritisation", "process improvement",
            "independent work", "adaptability",
        ],
        "experience": mytvs_experience([
            "Coordinate day-to-day HR administration, employee communication, and management follow-ups for a "
            "workforce of about 70 employees.",
            "Maintain confidential employee documents and organise information for accurate, timely retrieval.",
            "Prioritise recurring HR requirements independently, adapt to changing needs, and follow actions "
            "through to completion.",
        ]),
        "learning": ["performance", "excel", "uci", "vitara"],
        "technical": (
            "Microsoft Excel, Word, and PowerPoint · Claude, ChatGPT Codex, and Visual Studio Code for research, "
            "drafting, information organisation, and output review"
        ),
        "projects": [],
    },
    6: {
        "headline": "BUSINESS & MANAGEMENT | ECONOMICS | AI OUTPUT REVIEW",
        "summary": (
            "MHRM candidate and Economics graduate with approximately one year of HR Operations experience and "
            "a 12-week business-plan programme. Combines management principles, organisational behaviour, "
            "business operations, critical reading, structured reasoning, and AI-assisted output review. Scored "
            "90.4% in Performance Management and 78.8% in HR Analytics Using Excel."
        ),
        "skills": [
            "business management", "business operations", "management principles", "organisational behaviour",
            "business reasoning", "decision-making frameworks", "critical reading",
            "logical consistency", "structured explanation",
        ],
        "experience": mytvs_experience([
            "Coordinate day-to-day HR administration across a workforce of about 70, building practical context "
            "for people processes, organisational behaviour, and management follow-through.",
            "Organise employee documents and information with attention to accuracy, confidentiality, and timely "
            "retrieval.",
            "Prioritise recurring requirements independently and follow actions to completion in a lean "
            "business environment.",
        ]),
        "learning": ["performance", "excel", "uci", "vitara"],
        "technical": (
            "Claude and ChatGPT Codex for research, drafting, and output review · Visual Studio Code for "
            "structured information organisation · Microsoft Excel, Word, and PowerPoint"
        ),
        "projects": [copy.deepcopy(AI_WORKFLOW_PROJECT)],
    },
    7: {
        "headline": "CUSTOMER SUCCESS | DOCUMENTATION | EMPATHETIC SUPPORT",
        "summary": (
            "Early-career operations professional with approximately one year supporting a 70-person workforce "
            "through routine employee communication, documentation, and issue follow-up. Brings empathy, "
            "organisation, problem-solving, cross-functional awareness, and independent learning. Seeking to "
            "transfer this internal service foundation into technology customer success."
        ),
        "skills": [
            "empathetic communication", "employee support", "issue follow-up", "documentation",
            "problem-solving and critical thinking", "relationship awareness", "cross-functional teamwork",
            "technology interest", "independent learning",
        ],
        "experience": mytvs_experience([
            "Support a workforce of about 70 through routine employee communication, HR administration, and "
            "management follow-up, developing a service-oriented approach to internal stakeholders.",
            "Maintain clear, confidential employee documentation for accurate and timely retrieval.",
            "Own recurring priorities independently, coordinate management input, and follow actions through to "
            "completion.",
        ]),
        "learning": ["performance", "excel", "powerbi"],
        "technical": (
            "Microsoft Excel, Word, and PowerPoint · Claude, ChatGPT Codex, and Visual Studio Code for research, "
            "drafting, information organisation, and output review · introductory Power BI workshop"
        ),
        "projects": [],
    },
    8: {
        "headline": "BUSINESS DEVELOPMENT | RESEARCH | PERSISTENT FOLLOW-UP",
        "summary": (
            "Economics graduate and MHRM candidate with approximately one year of HR Operations experience and "
            "a 12-week business-plan programme. Brings business research, professional communication, structured "
            "follow-up, information synthesis, curiosity, and persistence. Interested in learning technology "
            "sales, prospect research, and Canonical's open-source portfolio."
        ),
        "skills": [
            "business research", "market research", "information synthesis", "professional communication",
            "structured follow-up", "economics", "business reasoning", "technology interest",
            "continuous learning",
        ],
        "experience": mytvs_experience([
            "Coordinate routine employee communication and management follow-ups across a workforce of about "
            "70, keeping outstanding actions moving to completion.",
            "Organise confidential employee information for accurate retrieval and structured follow-up.",
            "Work independently in a lean environment, prioritising changing requests and communicating clearly "
            "with management.",
        ]),
        "learning": ["excel", "performance", "powerbi"],
        "technical": (
            "Microsoft Excel, Word, and PowerPoint · Claude, ChatGPT Codex, and Visual Studio Code for business "
            "research, drafting, information organisation, and output review · introductory Power BI workshop"
        ),
        "projects": [],
    },
    9: {
        "headline": "ECONOMICS | MARKET REASONING | AI OUTPUT REVIEW",
        "summary": (
            "Economics graduate pursuing an MHRM, with approximately one year of HR Operations experience, a "
            "12-week business-plan programme, and analytics coursework. Brings macroeconomic reasoning, critical "
            "reading, logical analysis, careful interpretation of evidence, and hands-on AI output review. Ready "
            "to apply this foundation to entry-level investment reasoning and structured AI evaluation tasks."
        ),
        "skills": [
            "economics and macroeconomics", "investment fundamentals", "financial reasoning",
            "market reasoning", "performance measurement", "risk concepts", "analytical thinking",
            "critical reading", "logical consistency", "structured explanation",
        ],
        "experience": mytvs_experience([
            "Coordinate recurring HR administration for a workforce of about 70 with disciplined records "
            "handling, management follow-up, and action completion.",
            "Organise employee documentation and information with attention to accuracy, confidentiality, and "
            "timely retrieval.",
            "Prioritise recurring requirements independently and adapt to changing needs in a lean organisation.",
        ]),
        "learning": ["excel", "performance", "uci", "powerbi"],
        "technical": (
            "Claude and ChatGPT Codex for research, drafting, and output review · Visual Studio Code for "
            "structured information organisation · Microsoft Excel, Word, and PowerPoint · introductory Power BI"
        ),
        "projects": [copy.deepcopy(AI_WORKFLOW_PROJECT)],
    },
}

for _job in JOBS:
    _job.update(RECRUITER_REWRITES[_job["order"]])


FORMAT_RESEARCH = {
    "Ashby": (
        "Ashby parses uploaded resumes into candidate profiles and supports full-text resume search. A "
        "single-column, text-extractable PDF is the primary file; the DOCX is a fallback."
    ),
    "Workable": (
        "The live Workable form requires a resume and separately collects salary, availability, UK-hours, "
        "equipment, connectivity, and examples. Use the concise PDF and answer those fields separately."
    ),
    "Recruit CRM careers ATS": (
        "The official role-specific Recruit CRM application remains live. Use the exact 2026 India role title "
        "and a conventional one-column PDF; a DOCX fallback is included."
    ),
    "Canonical Careers": (
        "Canonical evaluates the CV alongside detailed application questions and, if advanced, an anonymous "
        "written interview. The CV therefore emphasises academics, initiative, evidence, and exact role fit; "
        "all written answers must be the candidate's own words."
    ),
    "Greenhouse": (
        "Greenhouse officially accepts DOC, DOCX, PDF, RTF, and TXT files up to 100 MB. Use the stable, "
        "text-extractable PDF and review the parsed application preview."
    ),
}


def filename_stem(job: dict) -> str:
    base = f"Mohammad_Azimuddin_{job['company']}_{job['role']}"
    return re.sub(r"[^A-Za-z0-9_-]+", "_", base).strip("_")


def set_cell_shading(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def set_spacing(paragraph, before=0, after=0, line=1.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_spacing(p, before=2.5, after=1.2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.3)
    run.font.color.rgb = RGBColor(23, 58, 86)
    set_cell_shading(p, "7791A4")


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=0.2, line=0.94)
    p.paragraph_format.left_indent = Cm(0.45)
    p.paragraph_format.first_line_indent = Cm(-0.22)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(8.3)


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 7.8) -> None:
    """Append a compact external hyperlink without relying on a template style."""
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "315F7E")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(round(size * 2)))
    properties.extend([fonts, color, underline, font_size])
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.set(qn("xml:space"), "preserve")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def write_docx(job: dict, path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.75)
    section.bottom_margin = Cm(0.70)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(0.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(8.5)
    normal.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=0.2)
    r = p.add_run(CONTACT["name"])
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor(23, 58, 86)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=0.1)
    r = p.add_run(job["headline"])
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(49, 95, 126)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=0.2)
    r = p.add_run(f"TARGET POSITION: {job['role'].upper()} · {job['company'].upper()}")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(7.6)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=0.8)
    contact = (
        f"{CONTACT['location']} · Contact via LinkedIn\n"
        f"{CONTACT['linkedin']}"
    )
    r = p.add_run(contact)
    r.font.name = "Arial"
    r.font.size = Pt(7.4)
    set_cell_shading(p, "254F70")

    add_section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    set_spacing(p, after=0.5, line=0.95)
    r = p.add_run(job["summary"])
    r.font.name = "Arial"
    r.font.size = Pt(8.6)

    add_section_heading(doc, "Core Skills")
    p = doc.add_paragraph()
    set_spacing(p, after=0.5, line=0.94)
    r = p.add_run(" · ".join(job["skills"]))
    r.font.name = "Arial"
    r.font.size = Pt(8.1)

    add_section_heading(doc, "Experience")
    for item in job["experience"]:
        p = doc.add_paragraph()
        set_spacing(p, before=0.5, after=0.1)
        r = p.add_run(f"{item['title']} | {item['dates']}")
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(8.8)
        p = doc.add_paragraph()
        set_spacing(p, after=0.1)
        r = p.add_run(item["org"])
        r.italic = True
        r.font.name = "Arial"
        r.font.size = Pt(7.9)
        r.font.color.rgb = RGBColor(66, 83, 103)
        for bullet in item["bullets"]:
            add_bullet(doc, bullet)

    if job.get("projects"):
        add_section_heading(doc, "Selected Project")
        for item in job["projects"]:
            p = doc.add_paragraph()
            set_spacing(p, before=0.5, after=0.1)
            r = p.add_run(f"{item['title']} | {item['dates']}")
            r.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(8.8)
            p = doc.add_paragraph()
            set_spacing(p, after=0.1)
            r = p.add_run(item["org"])
            r.italic = True
            r.font.name = "Arial"
            r.font.size = Pt(7.9)
            r.font.color.rgb = RGBColor(66, 83, 103)
            for bullet in item["bullets"]:
                add_bullet(doc, bullet)

    add_section_heading(doc, "Education")
    for item in EDUCATION:
        detail = f" · {item['detail']}" if item["detail"] else ""
        p = doc.add_paragraph()
        set_spacing(p, after=0.1)
        r = p.add_run(f"{item['title']} | {item['dates']}")
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r = p.add_run(f"\n{item['org']}{detail}")
        r.font.name = "Arial"
        r.font.size = Pt(7.9)
        r.font.color.rgb = RGBColor(66, 83, 103)

    add_section_heading(doc, "Certifications & Applied Learning")
    p = doc.add_paragraph()
    set_spacing(p, after=0.4, line=0.93)
    for index, key in enumerate(BASE.learning_keys(job)):
        if index:
            separator = p.add_run(" · ")
            separator.font.name = "Arial"
            separator.font.size = Pt(7.8)
        text_run = p.add_run(LEARNING[key])
        text_run.font.name = "Arial"
        text_run.font.size = Pt(7.8)
        if key == "oracle":
            add_hyperlink(p, " [Verify]", BASE.ORACLE_BADGE_URL, size=7.8)

    add_section_heading(doc, "Technical Skills")
    p = doc.add_paragraph()
    set_spacing(p, line=0.94)
    r = p.add_run(job["technical"])
    r.font.name = "Arial"
    r.font.size = Pt(8.0)

    props = doc.core_properties
    props.title = f"Mohammad Azimuddin — {job['company']} — {job['role']}"
    props.subject = "Role-specific ATS-friendly CV"
    props.author = "Mohammad Azimuddin"
    props.keywords = ", ".join(job["supported_keywords"])
    doc.save(path)


def resume_html(job: dict) -> str:
    """Create the accessible HTML reference used for text QA and manual review."""
    esc = BASE.esc

    def entries(items: list[dict]) -> str:
        blocks = []
        for item in items:
            bullets = "".join(f"<li>{esc(bullet)}</li>" for bullet in item["bullets"])
            blocks.append(
                f"""<div class="entry">
  <div class="entry-head"><strong>{esc(item['title'])} | {esc(item['dates'])}</strong></div>
  <div class="sub">{esc(item['org'])}</div>
  <ul>{bullets}</ul>
</div>"""
            )
        return "\n".join(blocks)

    education = []
    for item in EDUCATION:
        detail = f" · {esc(item['detail'])}" if item["detail"] else ""
        education.append(
            f"""<div class="entry compact">
  <div class="entry-head"><strong>{esc(item['title'])} | {esc(item['dates'])}</strong></div>
  <div class="sub">{esc(item['org'])}{detail}</div>
</div>"""
        )

    project_section = ""
    if job.get("projects"):
        project_section = f"<h2>SELECTED PROJECT</h2>\n{entries(job['projects'])}"

    skills = " · ".join(esc(skill) for skill in job["skills"])
    learning = "".join(
        (
            f'<li><a href="{esc(BASE.ORACLE_BADGE_URL)}">{esc(LEARNING[key])}</a> '
            f'<a href="{esc(BASE.ORACLE_BADGE_URL)}">[Verify]</a></li>'
            if key == "oracle"
            else f"<li>{esc(LEARNING[key])}</li>"
        )
        for key in BASE.learning_keys(job)
    )
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{esc(CONTACT['name'].title())} — {esc(job['company'])} — {esc(job['role'])}</title>
<style>
  @page {{ size: A4; margin: 9mm 12mm; }}
  * {{ box-sizing: border-box; }}
  body {{ margin: 0 auto; max-width: 186mm; color: #17202b; font: 9.1pt/1.24 Arial, Helvetica, sans-serif; }}
  .header {{ text-align: center; border-bottom: 1.8px solid #254f70; padding-bottom: 4px; }}
  h1 {{ margin: 0; color: #173a56; font-size: 20pt; letter-spacing: .35px; }}
  .headline {{ margin: 2px 0; color: #315f7e; font-size: 9pt; font-weight: 700; }}
  .target {{ margin: 1px 0 3px; font-size: 8.1pt; font-weight: 700; }}
  .contact {{ font-size: 8pt; }}
  a {{ color: inherit; text-decoration: none; }}
  h2 {{ margin: 5.5px 0 2.5px; padding-bottom: 1px; border-bottom: .7px solid #7791a4; color: #173a56; font-size: 10pt; letter-spacing: .2px; }}
  p {{ margin: 1.5px 0; }}
  ul {{ margin: 1.5px 0 2.5px 15px; padding: 0; }}
  li {{ margin: .8px 0; }}
  .skills {{ font-size: 8.7pt; line-height: 1.25; }}
  .entry {{ break-inside: avoid; margin-bottom: 2.5px; }}
  .entry.compact {{ margin-bottom: 1.5px; }}
  .entry-head {{ margin-bottom: .5px; }}
  .sub {{ color: #425367; font-size: 8.4pt; }}
  .technical {{ font-size: 8.6pt; }}
</style>
</head>
<body>
<div class="header">
  <h1>{esc(CONTACT['name'])}</h1>
  <div class="headline">{esc(job['headline'])}</div>
  <div class="target">TARGET POSITION: {esc(job['role'].upper())} · {esc(job['company'].upper())}</div>
  <div class="contact">
    {esc(CONTACT['location'])} · Contact via LinkedIn<br>
    <a href="{esc(CONTACT['linkedin_url'])}">{esc(CONTACT['linkedin'])}</a>
  </div>
</div>
<h2>PROFESSIONAL SUMMARY</h2>
<p>{esc(job['summary'])}</p>
<h2>CORE SKILLS</h2>
<p class="skills">{skills}</p>
<h2>EXPERIENCE</h2>
{entries(job['experience'])}
{project_section}
<h2>EDUCATION</h2>
{''.join(education)}
<h2>CERTIFICATIONS &amp; APPLIED LEARNING</h2>
<ul>{learning}</ul>
<h2>TECHNICAL SKILLS</h2>
<p class="technical">{esc(job['technical'])}</p>
</body>
</html>
"""


def application_notes(job: dict) -> str:
    supported = "\n".join(f"- {item}" for item in job["supported_keywords"])
    gaps = "\n".join(f"- {item}" for item in job["gap_keywords"])
    evidence_gaps = "\n".join(f"- {item}" for item in job["gaps"])
    format_note = FORMAT_RESEARCH[job["platform"]]
    project_check = (
        "- [ ] Be ready to explain the evidence-controlled application-system project, including your own "
        "workflow, the nine-pack output, and the checks you personally understand.\n"
        if job.get("projects") else ""
    )
    return f"""# {job['company']} — {job['role']}

Prepared and vacancy rechecked: 5 August 2026

## Application

- Official/live page: {job['url']}
- Platform: {job['platform']}
- Fit: {job['fit']}
- Status: {job['status']}

## CV and ATS pattern researched

{format_note}

The supplied CV is one A4 page, single-column, text-extractable, and uses conventional headings. Its opening
scan shows the exact target title, role evidence, supported scale, and strongest relevant learning. It has no
photo, icons, text boxes, columns, graphics, skill bars, date of birth, marital status, or unsupported metrics.

## Role pattern

{job['company_analysis']}

## Evidence-supported keywords included

{supported}

## Advertised keywords deliberately not claimed as experience

{gaps}

## Material evidence or eligibility gaps

{evidence_gaps}

## Suggested truthful positioning

{job['pitch']}

## Critical instruction

{job.get('special_instruction', 'Answer every application and screening question truthfully.')}

## Submission checklist

- [ ] Reopen the vacancy immediately before submitting.
- [ ] Upload the PDF from `00_Ready_to_Upload`; use DOCX only if the site has a PDF problem.
- [ ] Check the parsed name, phone, email, dates, education, and employer before continuing.
- [ ] Ensure signed myTVS evidence eventually supports `HR Operations · July 2025–Present`.
{project_check}- [ ] Be ready to explain every metric, tool, course, and experience bullet in your own words.
- [ ] Do not add HRIS, payroll, ATS, CRM, customer tickets, sales pipeline, Linux, advanced analytics,
      professional AI training, or quantified achievements without evidence.
- [ ] Retain a copy of the submitted CV and every application answer.

Targeting improves relevance, but no CV can guarantee a shortlist or offer.
"""


def master_readme() -> str:
    rows = []
    for job in JOBS:
        stem = filename_stem(job)
        rows.append(
            f"| {job['order']} | {job['company']} | {job['role']} | "
            f"[{stem}_Resume.pdf]({job['slug']}/{stem}_Resume.pdf) | "
            f"[{stem}_Resume.docx]({job['slug']}/{stem}_Resume.docx) | {job['status']} |"
        )
    return f"""# Current High-Paying Remote Applications — 5 August 2026

This pack contains nine separately targeted, senior-recruiter-reviewed, evidence-controlled CVs. Every role
has its own PDF, DOCX, plain-text copy, HTML reference, application notes, supported-keyword list, and gap
warning. The first scan leads with the exact role, supported scale, role evidence, and strongest learning;
experience bullets are rewritten for the employer rather than reused unchanged across all nine versions.

## CV index

| Priority | Company | Target role | PDF | DOCX | Status |
|---:|---|---|---|---|---|
{chr(10).join(rows)}

## Which file to upload

Use the PDF in `00_Ready_to_Upload` by default. These PDFs are one-page, A4, single-column and
text-extractable. DOCX copies are supplied as fallbacks for systems or recruiters that explicitly prefer Word.

## Evidence controls

- Current employment remains `HR Operations · myTVS · July 2025–Present` until a signed last working date exists.
- The myTVS wording is limited to candidate-confirmed HR administration, employee communication, management
  follow-up, documentation, confidentiality, independent prioritisation, and follow-through.
- MHRM is shown as 2024–2026 without claiming that a final degree has been issued.
- No HRIS, payroll, ATS, CRM, SaaS tickets, sales pipeline, Ubuntu/Linux, professional AI-training, advanced
  analytics, or invented metric has been added.
- Supported measurable proof points are limited to approximately one year of HR Operations, a workforce of
  about 70 employees, the seven-week HR internship, the 12-week SURE programme, and verified course scores.
- The evidence-controlled application-system project appears only where it supports the role; it must be
  discussed accurately as an AI-assisted independent project.

The CVs are tailored to improve relevance and parsing. They cannot guarantee a shortlist, interview, or offer.
"""


def write_keyword_matrix() -> None:
    with (OUT / "00_ATS_Keyword_Matrix.csv").open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(
            [
                "Priority", "Company", "Role", "Platform", "Supported keywords used",
                "Advertised keywords not claimed", "Fit", "Status", "Application URL",
            ]
        )
        for job in JOBS:
            writer.writerow(
                [
                    job["order"], job["company"], job["role"], job["platform"],
                    "; ".join(job["supported_keywords"]), "; ".join(job["gap_keywords"]),
                    job["fit"], job["status"], job["url"],
                ]
            )


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    ready = OUT / "00_Ready_to_Upload"
    ready.mkdir(parents=True, exist_ok=True)
    (OUT / "README.md").write_text(master_readme(), encoding="utf-8")
    write_keyword_matrix()

    # BASE.resume_html reads these globals from BASE; align the education/learning data explicitly.
    BASE.CONTACT = CONTACT
    BASE.EDUCATION = EDUCATION
    BASE.LEARNING = LEARNING
    RENDERER.OUT = OUT

    for job in JOBS:
        folder = OUT / job["slug"]
        folder.mkdir(parents=True, exist_ok=True)
        stem = filename_stem(job)

        html_source = resume_html(job)
        html_path = folder / f"{stem}_CV.html"
        txt_path = folder / f"{stem}_CV.txt"
        docx_path = folder / f"{stem}_Resume.docx"
        html_path.write_text(html_source, encoding="utf-8")
        txt_path.write_text(BASE.html_to_text(html_source), encoding="utf-8")
        write_docx(job, docx_path)
        (folder / "APPLICATION_NOTES.md").write_text(application_notes(job), encoding="utf-8")

        pdf_path = RENDERER.render_job(BASE, job)
        shutil.copy2(pdf_path, ready / f"{job['order']:02d}_{stem}_Resume.pdf")
        shutil.copy2(docx_path, ready / f"{job['order']:02d}_{stem}_Resume.docx")
        print(f"{job['order']:02d} {job['company']}: {pdf_path.name}")

    print(f"Generated {len(JOBS)} targeted application packages under {OUT}")


if __name__ == "__main__":
    main()
